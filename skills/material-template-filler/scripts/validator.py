#!/usr/bin/env python3
"""
格式校验器 - 检查填充后的文档是否符合要求
"""

from docx import Document
import re
import os


class Validator:
    def __init__(self, doc_path: str):
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"文档不存在：{doc_path}")
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.issues = []
        self.warnings = []
        
    def validate(self, template_structure: dict = None) -> dict:
        """校验文档"""
        self._check_empty_sections(template_structure)
        self._check_word_limits(template_structure)
        self._check_format()
        self._check_placeholders()
        
        return {
            'valid': len(self.issues) == 0 and len(self.warnings) == 0,
            'issues': self.issues,
            'warnings': self.warnings,
            'summary': self._get_summary()
        }
    
    def _check_empty_sections(self, template_structure: dict = None) -> None:
        """检查空模块"""
        if not template_structure:
            return
            
        filled_titles = set()
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                filled_titles.add(para.text.strip())
        
        for section in template_structure.get('sections', []):
            title = section['title']
            if title not in filled_titles:
                self.warnings.append(f"⚠️ 模块 '{title}' 未找到对应内容")
    
    def _check_word_limits(self, template_structure: dict = None) -> None:
        """检查字数限制"""
        if not template_structure:
            return
            
        current_section = None
        current_words = 0
        
        for para in self.doc.paragraphs:
            style = para.style.name
            
            if style.startswith('Heading'):
                # 检查上一节的字数
                if current_section and current_section.get('word_limit'):
                    limit = current_section['word_limit']
                    if current_words > limit:
                        self.warnings.append(
                            f"⚠️ '{current_section['title']}' 字数超限："
                            f"{current_words}/{limit}"
                        )
                    elif current_words < limit * 0.5:
                        self.warnings.append(
                            f"⚠️ '{current_section['title']}' 字数偏少："
                            f"{current_words}/{limit}"
                        )
                
                # 找到当前节
                for s in template_structure.get('sections', []):
                    if s['title'] == para.text.strip():
                        current_section = s
                        current_words = 0
                        break
            else:
                # 统计字数
                current_words += len(para.text.strip())
    
    def _check_format(self) -> None:
        """检查格式问题"""
        # 检查是否有明显的格式错误
        has_heading = False
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                has_heading = True
                break
        
        if not has_heading:
            self.warnings.append("⚠️ 文档中没有检测到标题样式")
    
    def _check_placeholders(self) -> None:
        """检查是否有未替换的占位符"""
        for para in self.doc.paragraphs:
            text = para.text
            if '{{' in text and '}}' in text:
                placeholders = re.findall(r'{{(\w+)}}', text)
                for ph in placeholders:
                    self.issues.append(f"❌ 发现未填充的占位符：{{{ph}}}")
    
    def _get_summary(self) -> str:
        """生成校验摘要"""
        lines = ["\n🔍 文档校验报告：\n"]
        
        if not self.issues and not self.warnings:
            lines.append("   ✅ 文档校验通过，无明显问题\n")
        else:
            if self.issues:
                lines.append("   问题：\n")
                for issue in self.issues:
                    lines.append(f"   - {issue}\n")
            
            if self.warnings:
                lines.append("\n   警告：\n")
                for warning in self.warnings:
                    lines.append(f"   - {warning}\n")
        
        return ''.join(lines)
    
    def print_report(self):
        """打印校验报告"""
        print(self._get_summary())


def main():
    import sys
    if len(sys.argv) < 2:
        print("用法：python validator.py <文档文件>")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    
    try:
        validator = Validator(doc_path)
        result = validator.validate()
        validator.print_report()
        
        print(f"\n📊 校验结果：{'通过' if result['valid'] else '需修正'}")
        
    except Exception as e:
        print(f"❌ 错误：{e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
