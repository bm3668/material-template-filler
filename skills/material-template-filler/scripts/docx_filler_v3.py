#!/usr/bin/env python3
"""
DocxFiller V3 - 基于 Anthropic docx skill 最佳实践改进

关键改进：
1. 智能语义匹配 - 自动分析章节与表格字段的对应关系
2. 表格字段精确匹配 - 直接匹配单元格标签
3. 内容净化 - 去除章节标题，只保留纯内容
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import os
import re

# 导入智能匹配器
from smart_matcher import SmartSectionMatcher, ContentFeatureAnalyzer

# 导入图片生成器
try:
    from image_generator import ImageGenerator, generate_mindmap_png
except ImportError:
    ImageGenerator = None
    generate_mindmap_png = None


class ContentPurifier:
    """
    内容净化器 - 从用户输入中提取纯内容，去除章节标题
    
    基于 Anthropic docx skill 的建议：编辑文档时直接操作内容，不要混入结构信息
    """
    
    # 章节标题模式
    SECTION_PATTERNS = [
        r'^[一二三四五六七八九十]+[、.．]',  # 一、二、三、
        r'^[（(][一二三四五六七八九十]+[)）]',  # （一）（二）
        r'^\d+[.．]',  # 1. 2. 3.
        r'^[\(（]\d+[\)）]',  # (1) (2)
        r'^第 [一二三四五六七八九十\d]+[章节]',  # 第一章 第 1 章
    ]
    
    def __init__(self, user_input: str):
        self.user_input = user_input
        self.sections = {}
        
    def extract_pure_content(self, section_title: str) -> str:
        """
        根据章节标题提取纯内容
        
        Args:
            section_title: 章节标题（如"摘要"、"设计目标"）
            
        Returns:
            净化后的纯内容（不含章节标题）
        """
        # 首先尝试直接匹配
        if section_title in self.sections:
            return self._purify_content(self.sections[section_title])
        
        # 模糊匹配
        for key, content in self.sections.items():
            if section_title.lower() in key.lower() or key.lower() in section_title.lower():
                return self._purify_content(content)
        
        return None
    
    def _purify_content(self, content: str) -> str:
        """
        净化内容 - 去除纯章节标题行，保留内容
        
        基于 Anthropic 建议：编辑时直接操作内容，不要混入结构
        
        注意：只过滤纯标题行（如"一、摘要"），保留带内容的行（如"1. 抗误码性能：..."）
        """
        if not content:
            return ""
        
        lines = content.split('\n')
        purified_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                purified_lines.append('')  # 保留空行
                continue
            
            # 检查是否是纯章节标题（没有实际内容）
            is_pure_header = False
            
            # 只过滤短的纯标题（长度<10 且匹配标题模式）
            if len(line) < 10:
                for pattern in self.SECTION_PATTERNS:
                    if re.match(pattern, line):
                        is_pure_header = True
                        break
            
            # 保留非纯标题行
            if not is_pure_header:
                purified_lines.append(line)
        
        return '\n'.join(purified_lines)
    
    def parse_input(self):
        """
        解析用户输入，提取各章节内容
        
        改进：支持 docx 文档的段落结构解析
        - 按行分割（而非空行），因为 docx 段落用\\n连接
        - 识别标题行后，收集后续非标题行作为该章节内容
        - 支持多级标题（一、/（一）/1. 等）
        - 聚合子章节内容到父章节
        """
        # 按行分割（支持 docx 段落结构）
        lines = self.user_input.split('\n')
        
        # 第一行特殊处理：可能是项目名称
        first_line = lines[0].strip() if lines else ''
        if first_line and len(first_line) < 100 and not any(re.match(p, first_line) for p in self.SECTION_PATTERNS):
            self.sections['项目名称'] = first_line
        
        current_section = None
        current_content = []
        section_hierarchy = []  # 记录章节层级 [(title, level, content_list), ...]
        
        for line in lines[1:]:  # 跳过第一行（已处理）
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是章节标题，并确定层级
            is_header = False
            header_title = None
            header_level = 3  # 默认最低层级
            
            for i, pattern in enumerate(self.SECTION_PATTERNS):
                match = re.match(pattern, line)
                if match:
                    header_title = line
                    header_level = i  # 模式索引越小，层级越高
                    is_header = True
                    break
            
            if is_header:
                # 保存之前的章节
                if current_section:
                    self.sections[current_section] = '\n'.join(current_content)
                    section_hierarchy.append((current_section, current_level, current_content))
                
                # 开始新章节
                current_section = header_title
                current_level = header_level
                current_content = []
            else:
                # 非标题行，添加到当前章节内容
                if current_section:
                    current_content.append(line)
        
        # 保存最后一个章节
        if current_section:
            self.sections[current_section] = '\n'.join(current_content)
            section_hierarchy.append((current_section, current_level, current_content))
        
        # 聚合子章节内容到父章节
        self._aggregate_section_content(section_hierarchy)
        
        # 如果没有解析到任何章节，尝试将整个输入作为"项目说明"
        if not self.sections and self.user_input.strip():
            remaining = '\n'.join(lines[1:]).strip() if lines else ''
            if remaining:
                self.sections['项目说明'] = remaining
        
        return self.sections
    
    def _aggregate_section_content(self, section_hierarchy: list):
        """
        聚合子章节内容到父章节
        
        例如："二、设计目标" 的内容可能为空，但其子章节"（一）技术目标"、"（二）非技术目标"有内容
        需要将子章节内容聚合到父章节，以便匹配
        
        改进：处理标题行本身包含内容的情况（如"1. 抗误码性能：设计实现..."）
        """
        # 为每个一级章节聚合其子章节内容
        parent_contents = {}
        current_parent = None
        
        for title, level, content in section_hierarchy:
            if level == 0:
                # 一级章节，初始化
                current_parent = title
                parent_contents[title] = list(content) if content else []
            elif current_parent:
                # 子章节，聚合到当前一级章节
                # 如果标题本身包含内容（如"1. 抗误码性能：..."），也添加到内容中
                title_content = self._extract_title_content(title, level)
                if title_content:
                    parent_contents[current_parent].append(title_content)
                if content:
                    parent_contents[current_parent].extend(content)
        
        # 更新 sections 字典
        for parent_title, aggregated_content in parent_contents.items():
            if aggregated_content:
                content_str = '\n'.join(aggregated_content)
                # 如果原来内容为空或很短，用聚合内容替换
                original = self.sections.get(parent_title, '')
                if len(original) < len(content_str):
                    self.sections[parent_title] = content_str
    
    def _extract_title_content(self, title: str, level: int) -> str:
        """
        从标题行中提取内容部分
        
        例如："1. 抗误码性能：设计实现 5G+ 低轨卫星异构信道下..." 
        应该返回 "设计实现 5G+ 低轨卫星异构信道下..."
        """
        if level == 0:
            # 一级标题通常不包含具体内容
            return ""
        
        if level == 1:
            # 二级标题：（一）技术目标（量化设计预期...）
            match = re.match(r'^[（(][一二三四五六七八九十]+[)）]\s*([^（(]*)', title)
            if match:
                content = match.group(1).strip()
                # 如果内容很短（<10 字），可能是纯标题
                if len(content) > 10:
                    return content
            return ""
        
        if level >= 2:
            # 三级及以下标题：1. 抗误码性能：设计实现...
            # 尝试提取冒号后的内容
            if '：' in title:
                parts = title.split('：', 1)
                if len(parts) == 2 and len(parts[1]) > 10:
                    return parts[1]
            elif ':' in title:
                parts = title.split(':', 1)
                if len(parts) == 2 and len(parts[1]) > 10:
                    return parts[1]
        
        return ""


class TableFieldMatcher:
    """
    表格字段匹配器 - 精确匹配表格中的字段标签
    
    基于 Anthropic docx skill：编辑表格时需要精确定位每个单元格
    
    使用智能匹配器结果进行内容查找
    """
    
    # 字段别名映射（用于直接匹配）
    FIELD_ALIASES = {
        '参赛赛项名称': ['参赛赛项', '赛项名称', '赛项'],
        '作品/方案名称': ['作品名称', '方案名称', '项目名称', '项目名'],
        '学校全称': ['学校名称', '学校', '大学', '学院'],
        '团队名称': ['团队', '队伍名称', '队名'],
        '联系电话': ['电话', '手机', '联系方式'],
        '邮箱': ['邮件', 'email', '电子邮箱'],
        '指导教师姓名': ['指导教师', '指导老师', '教师', '导师'],
        '队长姓名': ['队长', '负责人', '队长姓名'],
        '队员姓名': ['队员', '成员', '队员姓名'],
    }
    
    def __init__(self, content_map: dict, field_matches: dict = None):
        self.content_map = content_map
        self.field_matches = field_matches or {}
        
    def find_content(self, field_label: str) -> tuple:
        """
        根据字段标签查找匹配内容
        
        优先级：
        1. 智能匹配器结果（最高优先级）
        2. 直接匹配
        3. 别名匹配
        4. 模糊匹配
        
        Returns:
            (content, confidence) - 内容和置信度
        """
        # 1. 首先检查智能匹配器结果
        if self.field_matches and field_label in self.field_matches:
            section_title, confidence = self.field_matches[field_label]
            if section_title in self.content_map:
                content = self.content_map[section_title]
                return (self._purify_field_content(content), confidence)
        
        # 2. 直接匹配
        if field_label in self.content_map:
            content = self.content_map[field_label]
            return (self._purify_field_content(content), 0.9)
        
        # 3. 别名匹配
        for canonical, aliases in self.FIELD_ALIASES.items():
            if field_label == canonical or field_label in aliases:
                for alias in [canonical] + aliases:
                    if alias in self.content_map:
                        content = self.content_map[alias]
                        return (self._purify_field_content(content), 0.7)
        
        # 4. 模糊匹配
        for key, value in self.content_map.items():
            if field_label in key or key in field_label:
                return (self._purify_field_content(value), 0.5)
        
        return (None, 0)
    
    def _purify_field_content(self, content: str) -> str:
        """
        净化字段内容 - 去除章节标题
        
        例如：用户输入"一、摘要\n这是摘要内容..." -> 只返回"这是摘要内容..."
        """
        if not content:
            return ""
        
        lines = content.split('\n')
        purified = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 跳过纯章节标题行
            is_header = False
            for pattern in ContentPurifier.SECTION_PATTERNS:
                if re.match(pattern, line):
                    is_header = True
                    break
            
            if not is_header:
                purified.append(line)
        
        return '\n'.join(purified)


class DocxFillerV3:
    """
    专业级文档填充引擎 V3
    
    基于 Anthropic docx skill 最佳实践：
    - 精确的表格字段匹配
    - 内容净化（去除章节标题）
    - 表格样式优化
    """
    
    def __init__(self, template_path: str, output_path: str):
        self.template_path = template_path
        self.output_path = output_path
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在：{template_path}")
        
        shutil.copy(template_path, output_path)
        self.doc = Document(output_path)
        
        # 初始化
        self.table_matcher = None
        self.content_purifier = None
        self.image_generator = None
        
        # 填充报告
        self.fill_report = {
            'sections': {},
            'tables': [],
            'personal_info_missing': [],
            'images': [],
        }
        
        # 初始化图片生成器（如果可用）
        if ImageGenerator is not None:
            try:
                self.image_generator = ImageGenerator()
            except Exception as e:
                print(f"⚠️ 图片生成器初始化失败：{e}")
    
    def fill(self, user_input: str, content_map: dict = None) -> dict:
        """
        填充文档
        
        Args:
            user_input: 用户原始输入（用于内容净化）
            content_map: 已结构化的内容字典（可选）
        """
        # 初始化内容净化器
        self.content_purifier = ContentPurifier(user_input)
        self.content_purifier.parse_input()
        
        # 如果没有提供 content_map，使用净化后的内容
        if content_map is None:
            content_map = {}
            for section, content in self.content_purifier.sections.items():
                purified = self.content_purifier._purify_content(content)
                if purified:
                    content_map[section] = purified
        
        # 使用智能匹配器匹配章节到表格字段
        print("   智能分析章节与表格字段的对应关系...")
        self.smart_matcher = SmartSectionMatcher(content_map)
        self.field_matches = self.smart_matcher.match_all()
        
        # 打印匹配报告
        print("   " + self.smart_matcher.get_match_report().replace('\n', '\n   '))
        
        # 初始化表格匹配器（使用智能匹配结果）
        self.table_matcher = TableFieldMatcher(content_map, self.field_matches)
        
        # 填充段落
        self._fill_paragraphs(content_map)
        
        # 填充表格（重点改进）
        self._fill_tables_precise(content_map)
        
        # 保存
        self.doc.save(self.output_path)
        
        return self.fill_report
    
    def _fill_paragraphs(self, content_map: dict):
        """填充段落内容 - 改进版：支持替换占位符和匹配标题"""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else ''
            
            # 方法 1：识别标题并插入内容
            if style.startswith('Heading'):
                title = text
                if title in content_map:
                    content = content_map[title]
                    self._insert_after_heading(para, content)
                    self.fill_report['sections'][title] = {
                        'status': 'filled',
                        'word_count': len(content),
                        'confidence': 0.8
                    }
            
            # 方法 2：替换占位符文本（新增）
            # 检查是否包含常见占位符模式
            if text and len(text) < 200:  # 短段落可能是占位符
                for key, content in content_map.items():
                    # 检查段落是否包含与 key 相关的文字
                    if self._is_placeholder_for(text, key):
                        # 替换段落内容
                        self._replace_paragraph(para, content)
                        self.fill_report['sections'][key] = {
                            'status': 'replaced',
                            'word_count': len(content),
                            'confidence': 0.7
                        }
                        break
    
    def _is_placeholder_for(self, placeholder_text: str, section_key: str) -> bool:
        """判断占位符文本是否对应某个章节"""
        # 常见占位符关键词
        placeholder_keywords = ['请', '填写', '在此', 'describe', 'fill']
        
        # 如果包含占位符关键词
        has_placeholder = any(kw in placeholder_text for kw in placeholder_keywords)
        
        # 如果段落很短且包含占位符关键词，可能是对应章节的占位符
        if has_placeholder and len(placeholder_text) < 100:
            return True
        
        return False
    
    def _replace_paragraph(self, para, new_content: str):
        """替换段落内容"""
        # 清空原有内容
        para.clear()
        
        # 添加新内容
        run = para.add_run(new_content)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.insert(0, rFonts)
        
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)
    
    def _insert_after_heading(self, heading_para, content: str):
        """在标题后插入净化后的内容"""
        content_para = heading_para.insert_paragraph_before('')
        content_para.style = 'Normal'
        
        run = content_para.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.insert(0, rFonts)
        
        content_para.paragraph_format.line_spacing = 1.5
        content_para.paragraph_format.space_after = Pt(6)
    
    def _fill_tables_precise(self, content_map: dict):
        """
        精确填充表格
        
        基于 Anthropic docx skill 建议：
        - 精确定位每个单元格
        - 设置单元格边距
        - 保持表格样式
        
        模板结构识别：
        - 列 0：字段标签（如"摘要"、"设计目标"）
        - 列 1：填写说明（如"（作品/方案背景与意义...）"）或待填充内容
        - 内容应填充到列 1（右侧空格）
        """
        print("   处理表格字段...")
        
        # 个人信息字段
        personal_fields = [
            '学校全称', '学校名称', '团队名称', '联系电话', '邮箱',
            '指导教师', '指导教师姓名', '队长姓名', '队员姓名', '团队编号',
            '参赛赛项', '参赛赛项名称', '作品/方案名称',
        ]
        
        for table_idx, table in enumerate(self.doc.tables):
            print(f"     表格 {table_idx + 1}/{len(self.doc.tables)}")
            
            # 应用表格样式
            if table.style is None:
                table.style = 'Table Grid'
            
            # 遍历每一行
            for row_idx, row in enumerate(table.rows):
                # 跳过只有 1 列的行（标题行）
                if len(row.cells) < 2:
                    continue
                
                # 获取第一列作为字段标签
                label_cell = row.cells[0]
                label = label_cell.text.strip()
                
                # 跳过空标签或太长的标签（可能是内容）
                if not label or len(label) > 30:
                    continue
                
                # 跳过填写说明（以"（"开头）
                if label.startswith('（') or label.startswith('('):
                    continue
                
                # 查找匹配内容
                content, confidence = self.table_matcher.find_content(label)
                
                # 检查是否是个人信息字段
                is_personal = any(pf in label for pf in personal_fields)
                
                if content:
                    # 个人信息需要高置信度
                    if is_personal and confidence < 0.7:
                        print(f"       ⚠️ {label} -> 跳过（个人信息，置信度{confidence:.0%}）")
                        self.fill_report['personal_info_missing'].append({
                            'field': label,
                            'reason': '置信度低'
                        })
                        continue
                    
                    # 填充到右侧单元格（列 1）
                    target_cell = row.cells[1] if len(row.cells) > 1 else label_cell
                    self._fill_cell_precise(target_cell, content)
                    print(f"       ✅ {label} -> {len(content)}字 (填充到列 1)")
                    
                    self.fill_report['tables'].append({
                        'field': label,
                        'status': 'filled',
                        'word_count': len(content),
                        'confidence': confidence
                    })
                elif is_personal:
                    print(f"       ⚪ {label} -> 跳过（个人信息，用户未提供）")
                    self.fill_report['personal_info_missing'].append({
                        'field': label,
                        'reason': '用户未提供'
                    })
    
    def _fill_cell_precise(self, cell, content: str):
        """
        精确填充单元格
        
        基于 Anthropic docx skill：
        - 设置单元格边距
        - 保持格式
        """
        # 清空单元格
        cell.text = ''
        
        # 添加内容
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.insert(0, rFonts)
        
        # 设置段落边距（基于 Anthropic 建议）
        # 单元格内部填充，让内容不紧贴边框
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert(0, pBdr)
    
    def insert_image(self, image_path: str, caption: str = None, width: float = 15.0):
        """
        在文档末尾插入图片
        
        Args:
            image_path: 图片文件路径
            caption: 图片说明文字（可选）
            width: 图片宽度（厘米，默认 15cm）
            
        Returns:
            是否插入成功
        """
        if not os.path.exists(image_path):
            print(f"❌ 图片文件不存在：{image_path}")
            return False
        
        try:
            # 添加分页
            self.doc.add_page_break()
            
            # 插入图片
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 计算宽高比（如果是 PNG）
            height = None
            if image_path.endswith('.png'):
                try:
                    from PIL import Image
                    with Image.open(image_path) as img:
                        aspect_ratio = img.height / img.width
                        height = Cm(width * aspect_ratio)
                except:
                    height = Cm(width * 0.75)  # 默认宽高比
            
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width), height=height)
            
            # 添加说明文字
            if caption:
                caption_para = self.doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.font.size = Pt(9)
                caption_run.font.color.rgb = None  # 使用默认颜色（黑色）
                caption_run.italic = True
            
            self.fill_report['images'].append({
                'path': image_path,
                'caption': caption,
                'status': 'inserted'
            })
            
            print(f"✅ 图片已插入：{image_path}")
            return True
            
        except Exception as e:
            print(f"❌ 插入图片失败：{e}")
            return False
    
    def insert_mindmap(self, mermaid_content: str, caption: str = "思维导图", 
                      output_path: str = None, method: str = 'dot'):
        """
        生成并插入思维导图
        
        Args:
            mermaid_content: Mermaid 格式的思维导图定义
            caption: 图片说明文字
            output_path: 输出 PNG 路径（可选）
            method: 生成方法 ('dot' 或 'mermaid')
            
        Returns:
            是否插入成功
        """
        if self.image_generator is None:
            print("❌ 图片生成器不可用")
            return False
        
        # 生成临时文件名
        if output_path is None:
            import tempfile
            output_path = os.path.join(tempfile.gettempdir(), f'mindmap_{os.getpid()}.png')
        
        # 生成图片
        png_path = self.image_generator.generate_mindmap(
            mermaid_content, 
            output_path, 
            method
        )
        
        if png_path and os.path.exists(png_path):
            return self.insert_image(png_path, caption)
        else:
            print("❌ 思维导图生成失败")
            return False
    
    def generate_and_insert_diagram(self, text: str, diagram_type: str = 'mindmap',
                                   caption: str = None):
        """
        从文本生成图表并插入文档
        
        Args:
            text: 要分析的文本
            diagram_type: 图表类型 (flowchart, mindmap, pyramid)
            caption: 图片说明文字
            
        Returns:
            是否插入成功
        """
        if self.image_generator is None:
            print("❌ 图片生成器不可用")
            return False
        
        # 生成临时文件名
        import tempfile
        output_path = os.path.join(tempfile.gettempdir(), f'{diagram_type}_{os.getpid()}.png')
        
        # 从文本生成图表
        png_path = self.image_generator.generate_from_text(
            text, 
            diagram_type, 
            output_path
        )
        
        if png_path and os.path.exists(png_path):
            if caption is None:
                caption = f"{diagram_type.capitalize()} Diagram"
            return self.insert_image(png_path, caption)
        else:
            print("❌ 图表生成失败")
            return False
    
    def add_fill_report(self):
        """添加填充报告"""
        self.doc.add_page_break()
        
        # 模板可能没有 Heading 样式，使用普通段落加粗
        heading = self.doc.add_paragraph()
        run = heading.add_run('📋 自动填充报告')
        run.bold = True
        run.font.size = Pt(16)
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
        
        intro = self.doc.add_paragraph()
        intro.add_run('本文档由 材料模板填写 Skill 自动填充生成。\n').bold = True
        intro.add_run('请仔细检查以下内容。\n')
        
        self.doc.add_paragraph('─' * 50)
        
        # 表格字段填充情况
        if self.fill_report['tables']:
            heading = self.doc.add_paragraph()
            run = heading.add_run('表格字段填充情况')
            run.bold = True
            run.font.size = Pt(14)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(8)
            
            for item in self.fill_report['tables']:
                status = '✅' if item['confidence'] >= 0.7 else '⚠️'
                p = self.doc.add_paragraph()
                p.add_run(f'{status} {item["field"]}: {item["word_count"]}字 (置信度{item["confidence"]:.0%})')
        
        # 缺失的个人信息
        if self.fill_report['personal_info_missing']:
            heading = self.doc.add_paragraph()
            run = heading.add_run('⚠️ 需要手动填写的个人信息')
            run.bold = True
            run.font.size = Pt(14)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(8)
            
            for item in self.fill_report['personal_info_missing']:
                p = self.doc.add_paragraph()
                p.add_run(f'• {item["field"]}: {item["reason"]}')
        
        self.doc.save(self.output_path)


def main():
    """测试入口"""
    import sys
    
    if len(sys.argv) < 3:
        print("用法：python docx_filler_v3.py <模板文件> <输出文件>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # 测试输入
    test_input = """
星语智传——基于 MDVSC 的 5G+ 低轨卫星视频语义通信智能系统

一、摘要
当前 5G 广域覆盖存在短板，低轨卫星视频传输面临高误码、大时延、带宽稀缺的行业痛点。
本方案提出基于模型分割视频语义通信（MDVSC）的 5G+ 低轨卫星视频语义通信系统。

二、设计目标
1. 抗误码性能：1e-3 误码率时视频重建 PSNR≥32dB
2. 传输效率：带宽占用≤1.2Mbps（降低 70%）
3. 时延控制：端到端时延≤300ms

三、作品详情
本方案以"痛点聚焦 - 技术适配 - 架构简化 - 快速落地"为核心逻辑。
"""
    
    try:
        filler = DocxFillerV3(template_path, output_path)
        filler.fill(test_input)
        filler.add_fill_report()
        print(f"\n✅ 填充完成：{output_path}")
    except Exception as e:
        print(f"❌ 错误：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
