#!/usr/bin/env python3
"""
从 inputs 目录读取项目说明文档并填充模板

用法:
    python3 fill_from_inputs.py <模板路径> [项目说明文件路径]
    
如果未指定项目说明文件，自动使用 inputs/目录下最新的 .docx 文件
"""

import sys
import os
import re
from datetime import datetime

# 添加脚本目录到路径
script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)

from docx import Document
from template_parser import TemplateParser
from content_matcher import ContentMatcher
from docx_filler import DocxFiller
from report_generator import ReportGenerator


def extract_project_content(docx_path: str) -> dict:
    """
    从项目说明文档中提取结构化内容
    
    返回:
        dict: {
            '项目名称': str,
            '参赛赛项': str,
            '摘要': str,
            '设计目标': str,
            '作品详情': str,
            '经济社会价值': str,
            '进度计划': str,
            '学校名称': str,
            '团队名称': str,
            ...
        }
    """
    doc = Document(docx_path)
    
    # 初始化提取结果
    extracted = {
        '项目名称': '',
        '参赛赛项': '',
        '摘要': '',
        '设计目标': '',
        '作品详情': '',
        '经济社会价值': '',
        '进度计划': '',
        '学校名称': '',
        '团队名称': '',
        '联系电话': '',
        '邮箱': '',
        '指导教师': '',
        '队长姓名': '',
        '队员姓名': '',
    }
    
    # 章节标题映射（文档中的标题 → 提取字段）
    section_map = {
        '摘要': '摘要',
        '一、摘要': '摘要',
        '设计目标': '设计目标',
        '二、设计目标': '设计目标',
        '作品详情': '作品详情',
        '三、作品详情': '作品详情',
        '三、作品详情/解决方案详情': '作品详情',
        '解决方案详情': '作品详情',
        '经济与社会价值': '经济社会价值',
        '四、经济与社会价值': '经济社会价值',
        '经济社会价值': '经济社会价值',
        '项目进度计划': '进度计划',
        '五、项目进度计划': '进度计划',
        '五、项目进度计划（里程碑制定）': '进度计划',
        '进度计划': '进度计划',
    }
    
    # 解析文档内容
    current_section = None
    section_content = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是章节标题
        matched_section = None
        for title, field in section_map.items():
            if text.startswith(title):
                matched_section = field
                # 如果标题后有内容（如"摘要：xxx"），提取冒号后的内容
                if ':' in text or ':' in text:
                    parts = text.split(':', 1) if ':' in text else text.split(':', 1)
                    if len(parts) > 1 and len(parts[1].strip()) > 10:
                        section_content[field] = parts[1].strip()
                        matched_section = None  # 不再继续收集
                break
        
        if matched_section:
            current_section = matched_section
            if current_section not in section_content:
                section_content[current_section] = ''
        elif current_section:
            # 检查是否到了下一个大章节（如"六、"）
            if re.match(r'^[六七八九十][、. ]', text):
                current_section = None
            else:
                section_content[current_section] += text + '\n'
    
    # 尝试从文档标题提取项目名称
    if doc.paragraphs:
        first_para = doc.paragraphs[0].text.strip()
        if first_para and len(first_para) < 100:
            extracted['项目名称'] = first_para.replace('项目说明文档', '').strip(' ——')
    
    # 清理和赋值
    for field, content in section_content.items():
        if field in extracted:
            # 清理多余空白
            content = re.sub(r'\n+', '\n', content.strip())
            extracted[field] = content
    
    return extracted


def build_user_input(extracted: dict) -> str:
    """将提取的内容构建为用户输入格式"""
    parts = []
    
    if extracted.get('项目名称'):
        parts.append(f"项目名称：{extracted['项目名称']}")
    
    if extracted.get('参赛赛项'):
        parts.append(f"参赛赛项：{extracted['参赛赛项']}")
    
    if extracted.get('摘要'):
        # 摘要保持完整，通常不超过 500 字
        parts.append(f"摘要：{extracted['摘要']}")
    
    if extracted.get('设计目标'):
        # 设计目标保持完整，包含技术目标和非技术目标
        parts.append(f"设计目标：{extracted['设计目标']}")
    
    if extracted.get('作品详情'):
        # 作品详情保持完整，包含解决思路、可行性分析、技术方案、系统架构等
        parts.append(f"作品详情：{extracted['作品详情']}")
    
    if extracted.get('经济社会价值'):
        # 经济社会价值保持完整
        parts.append(f"经济社会价值：{extracted['经济社会价值']}")
    
    if extracted.get('进度计划'):
        # 进度计划保持完整，清理格式但不截断
        progress = extracted['进度计划']
        # 移除"五、项目进度计划"等标题前缀
        progress = re.sub(r'^五、项目进度计划 [（(].*?[)）]', '', progress).strip()
        # 移除"本项目以"之前的内容（如果有）
        match = re.search(r'本项目以.*', progress, re.DOTALL)
        if match:
            progress = match.group(0)
        # 清理可能混入的其他章节内容（如"3. 深化产教融合"）
        lines = progress.split('\n')
        clean_lines = []
        for line in lines:
            # 跳过明显是其他章节的内容
            if '深化产教融合' in line or '经济社会价值' in line:
                continue
            clean_lines.append(line)
        progress = '\n'.join(clean_lines)
        parts.append(f"进度计划：{progress}")
    
    return '\n'.join(parts)


def find_latest_project_doc(inputs_dir: str) -> str:
    """查找 inputs 目录下最新的项目说明文档"""
    latest_file = None
    latest_time = 0
    
    for filename in os.listdir(inputs_dir):
        if filename.endswith('.docx') and 'README' not in filename:
            filepath = os.path.join(inputs_dir, filename)
            mtime = os.path.getmtime(filepath)
            if mtime > latest_time:
                latest_time = mtime
                latest_file = filepath
    
    return latest_file


def generate_output_path(template_path: str) -> str:
    """生成输出文件路径"""
    template_name = os.path.basename(template_path)
    base, ext = os.path.splitext(template_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    workspace_dir = os.path.expanduser("~/.openclaw/workspace")
    filled_dir = os.path.join(workspace_dir, "filled")
    os.makedirs(filled_dir, exist_ok=True)
    
    return os.path.join(filled_dir, f"{base}_filled_{timestamp}.docx")


def main():
    if len(sys.argv) < 2:
        print("📋 从 inputs 目录填充模板")
        print("=" * 50)
        print("\n用法:")
        print("  python3 fill_from_inputs.py <模板路径> [项目说明文件路径]")
        print("\n示例:")
        print("  python3 fill_from_inputs.py templates/template.docx")
        print("  python3 fill_from_inputs.py templates/template.docx inputs/项目说明.docx")
        sys.exit(1)
    
    template_path = sys.argv[1]
    project_doc_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 如果是相对路径，转换为绝对路径
    if not os.path.isabs(template_path):
        workspace_dir = os.path.expanduser("~/.openclaw/workspace")
        template_path = os.path.join(workspace_dir, template_path)
    
    # 自动查找项目说明文档
    if not project_doc_path:
        inputs_dir = os.path.expanduser("~/.openclaw/workspace/inputs")
        project_doc_path = find_latest_project_doc(inputs_dir)
        if not project_doc_path:
            print("❌ 未在 inputs/目录下找到项目说明文档")
            sys.exit(1)
        print(f"📁 自动使用最新项目说明：{os.path.basename(project_doc_path)}")
    elif not os.path.isabs(project_doc_path):
        workspace_dir = os.path.expanduser("~/.openclaw/workspace")
        project_doc_path = os.path.join(workspace_dir, project_doc_path)
    
    print("\n📋 从 inputs 目录填充模板")
    print("=" * 50)
    
    # 步骤 1: 提取项目说明内容
    print(f"\n📥 步骤 1/6: 读取项目说明文档...")
    print(f"   文件：{project_doc_path}")
    
    try:
        extracted = extract_project_content(project_doc_path)
        print(f"   提取到 {sum(1 for v in extracted.values() if v)} 个字段")
    except Exception as e:
        print(f"❌ 读取失败：{e}")
        sys.exit(1)
    
    # 步骤 2: 构建用户输入
    print(f"\n🔧 步骤 2/6: 构建结构化输入...")
    user_input = build_user_input(extracted)
    print(f"   输入长度：{len(user_input)} 字符")
    
    # 步骤 3: 解析模板
    print(f"\n🔍 步骤 3/6: 解析模板...")
    print(f"   模板：{template_path}")
    
    try:
        parser = TemplateParser(template_path)
        parser.parse()
        template_structure = parser.get_structure()
        print(f"   模板模块数：{len(template_structure['sections'])}")
    except Exception as e:
        print(f"❌ 模板解析失败：{e}")
        sys.exit(1)
    
    # 步骤 4: 解析表格字段要求
    print(f"\n📋 步骤 4/6: 解析表格字段要求...")
    
    try:
        from table_parser import TableFieldParser
        parser_obj = TemplateParser(template_path)
        parser_obj.parse()
        
        field_requirements = {}
        for table in parser_obj.doc.tables:
            field_parser = TableFieldParser()
            fields = field_parser.parse_table(table)
            for field in fields:
                std_name = field['std_name']
                if std_name not in field_requirements:
                    field_requirements[std_name] = {
                        'requirements': field['requirements'],
                        'word_limit': field['word_limit'],
                        'content_points': field['content_points'],
                    }
        
        print(f"   解析到 {len(field_requirements)} 个字段要求")
    except Exception as e:
        print(f"   ⚠️ 解析字段要求失败：{e}")
        field_requirements = {}
    
    # 步骤 5: 匹配内容
    print(f"\n🤖 步骤 5/6: 匹配内容到模板模块...")
    
    matcher = ContentMatcher(user_input, template_structure['sections'])
    matcher._extract_table_fields(field_requirements)
    match_result = matcher.match()
    
    # 步骤 6: 填充文档
    output_path = generate_output_path(template_path)
    print(f"\n✏️  步骤 6/6: 填充文档...")
    print(f"   输出：{output_path}")
    
    try:
        filler = DocxFiller(template_path, output_path)
        filler.fill(match_result['matches'])
    except Exception as e:
        print(f"❌ 文档填充失败：{e}")
        sys.exit(1)
    
    # 生成报告
    print(f"\n📄 生成填充报告...")
    
    try:
        report_generator = ReportGenerator(template_structure, match_result, user_input)
        report_path = report_generator.generate_report_path(template_path)
        report_generator.generate_report(report_path)
        print(f"   报告：{report_path}")
    except Exception as e:
        print(f"❌ 报告生成失败：{e}")
        report_path = None
    
    # 输出总结
    print("\n" + "=" * 50)
    print("✅ 完成！")
    print(f"\n📄 输出文件：{output_path}")
    if report_path:
        print(f"📊 填充报告：{report_path}")
    print(f"\n📊 填充置信度:")
    
    for title, conf in match_result['confidence'].items():
        if conf >= 0.7:
            status = "✅"
        elif conf >= 0.4:
            status = "⚠️"
        else:
            status = "❌"
        print(f"   {status} {title}: {conf:.0%}")
    
    print("\n💡 建议：打开文档检查填充内容，特别是 ⚠️ 和 ❌ 标记的部分。")
    print("=" * 50 + "\n")
    
    return output_path, report_path


if __name__ == "__main__":
    main()
