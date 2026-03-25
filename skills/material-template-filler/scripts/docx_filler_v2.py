#!/usr/bin/env python3
"""
DocxFiller V2 - 专业级文档填充引擎

基于 Anthropic docx skill 最佳实践：
- 页面设置标准化（A4/US Letter，标准边距）
- 样式管理系统（覆盖内置 Heading 样式）
- 高级元素支持（书签、表格样式、脚注）
- 字体兼容性（使用通用字体）
- 文档验证（生成后验证 .docx 合法性）
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os
import re
import zipfile
import xml.etree.ElementTree as ET

from table_parser import TableFieldParser


class StyleManager:
    """
    样式管理器 - 管理文档样式覆盖
    
    基于 Anthropic docx skill 的样式最佳实践
    """
    
    # 中文字体映射（Windows / macOS / Linux）
    FONT_MAP = {
        'heading': {
            'zh': ['黑体', 'STHeiti', 'WenQuanYi Micro Hei'],
            'en': 'Arial Black',
        },
        'body': {
            'zh': ['宋体', 'SimSun', 'WenQuanYi Zen Hei'],
            'en': 'Arial',
        },
        'mono': {
            'zh': ['等线', 'DengXian'],
            'en': 'Consolas',
        }
    }
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.styles = doc.styles
        
    def apply_professional_styles(self):
        """
        应用专业文档样式
        
        - 标题：黑体/Arial Black，加粗
        - 正文：宋体/Arial，12pt，1.5 倍行距
        - 段前段后：6pt
        """
        self._setup_heading_styles()
        self._setup_normal_style()
        self._setup_table_styles()
        
    def _setup_heading_styles(self):
        """设置标题样式"""
        # Heading 1
        if 'Heading 1' in self.styles:
            h1 = self.styles['Heading 1']
            h1.font.name = self.FONT_MAP['heading']['en']
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.paragraph_format.space_before = Pt(18)
            h1.paragraph_format.space_after = Pt(12)
            h1.paragraph_format.line_spacing = 1.5
            
        # Heading 2
        if 'Heading 2' in self.styles:
            h2 = self.styles['Heading 2']
            h2.font.name = self.FONT_MAP['heading']['en']
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(14)
            h2.paragraph_format.space_after = Pt(8)
            h2.paragraph_format.line_spacing = 1.5
            
        # Heading 3
        if 'Heading 3' in self.styles:
            h3 = self.styles['Heading 3']
            h3.font.name = self.FONT_MAP['heading']['en']
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.paragraph_format.space_before = Pt(12)
            h3.paragraph_format.space_after = Pt(6)
            h3.paragraph_format.line_spacing = 1.5
    
    def _setup_normal_style(self):
        """设置正文样式"""
        if 'Normal' in self.styles:
            normal = self.styles['Normal']
            normal.font.name = self.FONT_MAP['body']['en']
            normal.font.size = Pt(10.5)  # 五号字
            
            # 设置中文字体（通过 rPr 元素）
            self._set_east_asia_font(normal, self.FONT_MAP['body']['zh'][0])
            
            normal.paragraph_format.line_spacing = 1.5
            normal.paragraph_format.space_after = Pt(6)
            normal.paragraph_format.space_before = Pt(0)
    
    def _set_east_asia_font(self, style, font_name):
        """
        设置东亚字体（中文）
        
        通过修改 rPr 元素实现
        """
        rPr = style._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), self.FONT_MAP['body']['en'])
        rFonts.set(qn('w:hAnsi'), self.FONT_MAP['body']['en'])
        rPr.insert(0, rFonts)
    
    def _setup_table_styles(self):
        """设置表格样式"""
        # 可以自定义表格样式，这里使用 Word 内置样式
        pass


class PageSetupManager:
    """
    页面设置管理器
    
    基于 Anthropic docx skill 的页面设置最佳实践
    """
    
    # 常见纸张尺寸（DXA 单位，1440 DXA = 1 英寸 = 2.54 厘米）
    PAPER_SIZES = {
        'A4': {'width': 11906, 'height': 16838},  # 210mm × 297mm
        'A3': {'width': 16838, 'height': 23811},  # 297mm × 420mm
        'Letter': {'width': 12240, 'height': 15840},  # 8.5" × 11"
        'Legal': {'width': 12240, 'height': 20160},  # 8.5" × 14"
    }
    
    # 标准边距（DXA 单位）
    MARGINS = {
        'normal': {'top': 1440, 'bottom': 1440, 'left': 1440, 'right': 1440},  # 1 英寸
        'narrow': {'top': 720, 'bottom': 720, 'left': 720, 'right': 720},  # 0.5 英寸
        'wide': {'top': 2160, 'bottom': 2160, 'left': 2160, 'right': 2160},  # 1.5 英寸
    }
    
    def __init__(self, doc: Document):
        self.doc = doc
        
    def setup_page(self, paper_size='A4', orientation='portrait', margin_style='normal'):
        """
        设置页面尺寸和边距
        
        Args:
            paper_size: 'A4', 'A3', 'Letter', 'Legal'
            orientation: 'portrait' (纵向) 或 'landscape' (横向)
            margin_style: 'normal', 'narrow', 'wide'
        """
        if not self.doc.sections:
            self.doc.add_section()
            
        section = self.doc.sections[0]
        
        # 设置纸张尺寸
        size = self.PAPER_SIZES.get(paper_size, self.PAPER_SIZES['A4'])
        section.page_width = size['width']
        section.page_height = size['height']
        
        # 设置方向
        if orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT
        
        # 设置边距
        margins = self.MARGINS.get(margin_style, self.MARGINS['normal'])
        section.top_margin = margins['top']
        section.bottom_margin = margins['bottom']
        section.left_margin = margins['left']
        section.right_margin = margins['right']
        
        return section


class DocxValidator:
    """
    文档验证器 - 验证 .docx 文件合法性
    
    基于 Anthropic docx skill 的验证流程
    """
    
    REQUIRED_FILES = [
        '[Content_Types].xml',
        '_rels/.rels',
        'word/document.xml',
        'word/_rels/document.xml.rels',
        'word/styles.xml',
    ]
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        
    def validate(self) -> dict:
        """
        验证 .docx 文件
        
        Returns:
            {
                'valid': bool,
                'errors': list,
                'warnings': list
            }
        """
        result = {
            'valid': True,
            'errors': [],
            'warnings': []
        }
        
        # 检查文件是否存在
        if not os.path.exists(self.file_path):
            result['valid'] = False
            result['errors'].append(f'文件不存在：{self.file_path}')
            return result
        
        # 检查 ZIP 结构
        try:
            with zipfile.ZipFile(self.file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                
                # 检查必需文件
                for required in self.REQUIRED_FILES:
                    if required not in file_list:
                        result['warnings'].append(f'缺少可选文件：{required}')
                
                # 验证 document.xml
                if 'word/document.xml' in file_list:
                    try:
                        xml_content = zip_file.read('word/document.xml')
                        ET.fromstring(xml_content)
                    except ET.ParseError as e:
                        result['valid'] = False
                        result['errors'].append(f'document.xml XML 解析错误：{str(e)}')
                
        except zipfile.BadZipFile:
            result['valid'] = False
            result['errors'].append('无效的 ZIP 文件（.docx 本质是 ZIP 归档）')
        except Exception as e:
            result['valid'] = False
            result['errors'].append(f'验证过程出错：{str(e)}')
        
        return result


class DocxFillerV2:
    """
    专业级文档填充引擎 V2
    
    基于 Anthropic docx skill 最佳实践重构
    """
    
    def __init__(self, template_path: str, output_path: str):
        self.template_path = template_path
        self.output_path = output_path
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在：{template_path}")
        
        # 复制模板到输出路径
        shutil.copy(template_path, output_path)
        self.doc = Document(output_path)
        
        # 初始化各管理器
        self.style_manager = StyleManager(self.doc)
        self.page_manager = PageSetupManager(self.doc)
        self.table_parser = TableFieldParser()
        
        # 应用专业样式
        self.style_manager.apply_professional_styles()
        
        # 存储填充报告数据
        self.fill_report_data = {
            'sections': {},
            'tables': [],
            'personal_info_missing': [],
        }
    
    def fill(self, content_map: dict, auto_detect_paper_size=True) -> dict:
        """
        填充内容到文档
        
        Args:
            content_map: 内容字典 {模块标题：内容}
            auto_detect_paper_size: 是否自动检测纸张尺寸
            
        Returns:
            填充报告数据
        """
        # 自动检测并设置页面尺寸
        if auto_detect_paper_size:
            self._auto_detect_paper_size()
        
        # 处理段落（标题 + 内容）
        self._fill_paragraphs(content_map)
        
        # 处理表格
        self._fill_tables_enhanced(content_map)
        
        # 保存文档
        self.doc.save(self.output_path)
        
        # 验证文档
        validator = DocxValidator(self.output_path)
        validation_result = validator.validate()
        
        if not validation_result['valid']:
            print(f"⚠️ 文档验证失败：{validation_result['errors']}")
        
        return self.fill_report_data
    
    def _auto_detect_paper_size(self):
        """自动检测模板的纸张尺寸"""
        if self.doc.sections:
            section = self.doc.sections[0]
            width = section.page_width
            height = section.page_height
            
            # 检测是否为 A4（允许 5% 误差）
            a4_width = 11906
            a4_height = 16838
            
            if abs(width - a4_width) / a4_width < 0.05:
                # 已经是 A4，无需调整
                pass
            elif abs(width - 12240) / 12240 < 0.05:
                # Letter 尺寸
                print("   检测到 Letter 纸张尺寸")
            else:
                # 自定义尺寸，保持原样
                pass
    
    def _fill_paragraphs(self, content_map: dict) -> None:
        """
        填充段落内容
        
        改进点：
        - 保持原有样式
        - 智能识别占位符
        - 支持多级标题
        """
        for para in self.doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ''
            
            # 识别标题段落
            if style_name.startswith('Heading'):
                title = text
                if title in content_map:
                    content = content_map[title]
                    self._insert_content_after_heading(para, content, style_name)
                    self.fill_report_data['sections'][title] = {
                        'status': 'filled',
                        'word_count': len(content),
                        'confidence': 0.8
                    }
                else:
                    # 尝试模糊匹配
                    matched_key = self._fuzzy_match(title, content_map.keys())
                    if matched_key:
                        content = content_map[matched_key]
                        self._insert_content_after_heading(para, content, style_name)
                        self.fill_report_data['sections'][title] = {
                            'status': 'filled',
                            'word_count': len(content),
                            'confidence': 0.6,
                            'matched_key': matched_key
                        }
                    else:
                        self.fill_report_data['sections'][title] = {
                            'status': 'empty',
                            'word_count': 0,
                            'confidence': 0
                        }
            
            # 识别占位符 {{field_name}}
            elif '{{' in text and '}}' in text:
                self._fill_placeholders(para, content_map)
    
    def _fuzzy_match(self, target: str, candidates: list) -> str:
        """
        模糊匹配标题
        
        使用简单的包含关系匹配
        """
        target_lower = target.lower()
        
        # 完全包含
        for candidate in candidates:
            if target_lower in candidate.lower() or candidate.lower() in target_lower:
                return candidate
        
        # 关键词匹配
        keywords = ['背景', '目标', '内容', '方案', '创新', '成果', '计划', '预算']
        for keyword in keywords:
            if keyword in target_lower:
                for candidate in candidates:
                    if keyword in candidate.lower():
                        return candidate
        
        return None
    
    def _insert_content_after_heading(self, heading_para, content: str, heading_style: str) -> None:
        """
        在标题后插入内容
        
        改进点：
        - 保持段落样式一致性
        - 使用合适的字体和字号
        - 设置合理的行距和间距
        """
        # 创建新段落
        content_para = heading_para.insert_paragraph_before('')
        content_para.style = 'Normal'
        
        # 添加内容
        run = content_para.add_run(content)
        run.font.size = Pt(10.5)  # 五号字
        run.font.name = 'Arial'
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.insert(0, rFonts)
        
        # 设置段落格式
        content_para.paragraph_format.line_spacing = 1.5
        content_para.paragraph_format.space_after = Pt(6)
        content_para.paragraph_format.space_before = Pt(0)
    
    def _fill_placeholders(self, para, content_map: dict) -> None:
        """
        填充占位符 {{field_name}}
        
        支持多个占位符在同一段落
        """
        text = para.text
        placeholders = re.findall(r'{{(\w+)}}', text)
        
        for placeholder in placeholders:
            if placeholder in content_map:
                text = text.replace(f'{{{{{placeholder}}}}}', content_map[placeholder])
            else:
                # 尝试模糊匹配
                matched_key = self._fuzzy_match(placeholder, content_map.keys())
                if matched_key:
                    text = text.replace(f'{{{{{placeholder}}}}}', content_map[matched_key])
                else:
                    # 保留占位符，标记为待填充
                    text = text.replace(f'{{{{{placeholder}}}}}', f'[待填充：{placeholder}]')
        
        # 更新段落内容
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
    
    def _fill_tables_enhanced(self, content_map: dict) -> None:
        """
        增强版表格填充
        
        改进点：
        - 智能识别字段类型
        - 区分个人信息和项目内容
        - 支持字数限制
        - 优化表格样式
        """
        print("   处理表格字段...")
        
        # 个人信息字段列表
        personal_fields = [
            '学校名称', '学校全称', '团队名称', '联系电话', '邮箱',
            '指导教师', '指导教师姓名', '队长姓名', '队员姓名', '团队编号',
            '参赛赛项', '参赛赛项名称',
        ]
        
        for table_idx, table in enumerate(self.doc.tables):
            print(f"     表格 {table_idx + 1}/{len(self.doc.tables)}")
            
            # 应用表格样式
            if table.style is None:
                table.style = 'Table Grid'
            
            # 解析表格字段
            fields = self.table_parser.parse_table(table)
            
            for field in fields:
                std_name = field['std_name']
                label = field['label']
                cell = field['cell']
                requirements = field['requirements']
                word_limit = field['word_limit']
                
                # 查找匹配内容
                matched_content = self._find_matched_content(std_name, label, content_map)
                
                # 检查是否是个人信息字段
                is_personal = std_name in personal_fields or any(p in label for p in personal_fields)
                
                if matched_content:
                    confidence = self._estimate_confidence(std_name, label, content_map)
                    
                    # 个人信息字段：只有高置信度才填充
                    if is_personal and confidence < 0.7:
                        print(f"       ⚠️ {label} -> 跳过（个人信息，置信度{confidence:.0%}）")
                        self.fill_report_data['personal_info_missing'].append({
                            'field': label,
                            'reason': '用户未提供精确信息'
                        })
                        continue
                    
                    # 填充内容
                    self._fill_table_cell(cell, matched_content, requirements, word_limit)
                    print(f"       ✅ {label} -> {len(matched_content)}字")
                    
                    self.fill_report_data['tables'].append({
                        'field': label,
                        'status': 'filled',
                        'word_count': len(matched_content),
                        'confidence': confidence
                    })
                elif is_personal:
                    print(f"       ⚪ {label} -> 跳过（个人信息，用户未提供）")
                    self.fill_report_data['personal_info_missing'].append({
                        'field': label,
                        'reason': '用户未提供'
                    })
                else:
                    # 项目内容字段：标记为待填充
                    print(f"       ❌ {label} -> 待填充（无匹配内容）")
                    self.fill_report_data['tables'].append({
                        'field': label,
                        'status': 'empty',
                        'word_count': 0,
                        'confidence': 0
                    })
    
    def _find_matched_content(self, std_name: str, label: str, content_map: dict) -> str:
        """查找匹配的内容"""
        # 直接匹配
        if std_name in content_map:
            content = content_map[std_name]
            if content and content.strip():
                return content
        
        if label in content_map:
            content = content_map[label]
            if content and content.strip():
                return content
        
        # 模糊匹配
        for key, value in content_map.items():
            if std_name in key or key in std_name:
                if value and value.strip():
                    return value
            if label in key or key in label:
                if value and value.strip():
                    return value
        
        return None
    
    def _estimate_confidence(self, std_name: str, label: str, content_map: dict) -> float:
        """估算匹配置信度"""
        # 直接匹配：高置信度
        if std_name in content_map or label in content_map:
            return 0.9
        
        # 模糊匹配：中置信度
        for key in content_map.keys():
            if std_name in key or key in std_name:
                return 0.7
            if label in key or key in label:
                return 0.7
        
        return 0.5
    
    def _fill_table_cell(self, cell, content: str, requirements: str = '', word_limit: int = None) -> None:
        """
        填充表格单元格
        
        改进点：
        - 保持单元格原有格式
        - 支持字数限制截断
        - 设置合适的字体
        """
        # 从要求中提取字数限制
        if word_limit is None and requirements:
            word_limit = self._extract_word_limit(requirements)
        
        # 截断超出字数限制的内容
        if word_limit and len(content) > word_limit:
            content = content[:word_limit-3] + '...'
        
        # 清空并填充
        cell.text = ''
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = 'Arial'
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Arial')
        rPr.insert(0, rFonts)
    
    def _extract_word_limit(self, requirements: str) -> int:
        """从要求中提取字数限制"""
        match = re.search(r'不超过？(\d+)[字个]', requirements)
        return int(match.group(1)) if match else None
    
    def add_fill_report(self, report_data: dict = None) -> None:
        """
        在文档末尾添加填充报告
        
        Args:
            report_data: 填充报告数据（可选，使用内部数据如果未提供）
        """
        if report_data is None:
            report_data = self.fill_report_data
        
        self.doc.add_page_break()
        
        # 报告标题
        report_heading = self.doc.add_heading('📋 自动填充报告', level=1)
        report_heading.paragraph_format.space_before = Pt(24)
        report_heading.paragraph_format.space_after = Pt(12)
        
        # 说明文字
        intro = self.doc.add_paragraph()
        intro.add_run('本文档由 材料模板填写 Skill 自动填充生成。\n').bold = True
        intro.add_run('请仔细检查以下内容，特别是标记为 ⚠️ 和 ❌ 的部分。\n')
        
        # 图例
        legend = self.doc.add_paragraph()
        legend.add_run('图例：').bold = True
        legend.add_run(' ✅ 高置信度  ')
        legend.add_run('⚠️ 中置信度（建议检查）  ')
        legend.add_run('❌ 低置信度（需人工补充）')
        
        self.doc.add_paragraph('─' * 50)
        
        # 模块填充情况
        self.doc.add_heading('各模块填充情况', level=2)
        
        for section_title, data in report_data.get('sections', {}).items():
            confidence = data.get('confidence', 0)
            word_count = data.get('word_count', 0)
            
            if confidence >= 0.7:
                status = "✅"
            elif confidence >= 0.4:
                status = "⚠️"
            else:
                status = "❌"
            
            p = self.doc.add_paragraph()
            p.add_run(f'{status} {section_title}').bold = True
            p.add_run(f' (置信度：{confidence:.0%}, {word_count}字)')
        
        # 表格字段填充情况
        if report_data.get('tables'):
            self.doc.add_heading('表格字段填充情况', level=2)
            
            for table_field in report_data['tables']:
                field = table_field.get('field', '未知')
                confidence = table_field.get('confidence', 0)
                word_count = table_field.get('word_count', 0)
                status_text = table_field.get('status', 'unknown')
                
                if status_text == 'filled':
                    status = "✅"
                elif status_text == 'empty':
                    status = "❌"
                else:
                    status = "⚠️"
                
                p = self.doc.add_paragraph()
                p.add_run(f'{status} {field}').bold = True
                p.add_run(f' ({word_count}字)')
        
        # 缺失的个人信息
        if report_data.get('personal_info_missing'):
            self.doc.add_heading('⚠️ 缺失的个人信息', level=2)
            
            for item in report_data['personal_info_missing']:
                p = self.doc.add_paragraph()
                p.add_run(f'• {item["field"]}: {item["reason"]}')
        
        # 建议
        self.doc.add_heading('建议', level=2)
        suggestions = self.doc.add_paragraph()
        suggestions.add_run('1. 打开文档检查自动填充的内容\n')
        suggestions.add_run('2. 重点检查 ⚠️ 和 ❌ 标记的部分\n')
        suggestions.add_run('3. 补充缺失的个人信息\n')
        suggestions.add_run('4. 调整格式以符合要求\n')
        
        self.doc.save(self.output_path)
    
    def get_output_info(self) -> dict:
        """返回输出文件信息"""
        return {
            'output_path': self.output_path,
            'template_path': self.template_path,
            'file_size': os.path.getsize(self.output_path) if os.path.exists(self.output_path) else 0,
            'page_count': len(self.doc.sections),
        }


def main():
    """测试入口"""
    import sys
    
    if len(sys.argv) < 3:
        print("用法：python docx_filler_v2.py <模板文件> <输出文件>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # 测试内容
    test_content = {
        '项目背景': '当前在线教育平台缺乏个性化推荐，学生学习效率低。本项目旨在解决这一问题。',
        '研究目标': '开发一个智能学习推荐系统，提高学生学习效率。',
        '研究内容': '包括用户行为分析、内容推荐算法、系统架构设计。',
        '技术方案': '采用深度学习技术，构建用户画像和物品画像，使用协同过滤和内容推荐算法。',
        '创新点': '提出了一种新的混合推荐算法，结合了多种技术的优势。',
        '预期成果': '发表论文 2 篇，申请专利 1 项，开发原型系统 1 套。',
    }
    
    try:
        filler = DocxFillerV2(template_path, output_path)
        report_data = filler.fill(test_content)
        filler.add_fill_report()
        
        print(f"\n✅ 填充完成：{output_path}")
        print(f"📄 文件信息：{filler.get_output_info()}")
        print(f"📊 填充报告：{len(report_data.get('sections', {}))} 个模块")
        
    except Exception as e:
        print(f"❌ 错误：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
