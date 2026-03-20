#!/usr/bin/env python3
"""
docx 填充器 - 将匹配内容写入模板（增强版，支持表格字段智能填充）
"""

from docx import Document
from docx.shared import Pt, Inches
import shutil
import os
import re

from table_parser import TableFieldParser


class DocxFiller:
    def __init__(self, template_path: str, output_path: str):
        self.template_path = template_path
        self.output_path = output_path
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在：{template_path}")
        
        # 先复制模板，避免修改原文件
        shutil.copy(template_path, output_path)
        self.doc = Document(output_path)
        
        # 初始化表格字段解析器
        self.table_parser = TableFieldParser()
        
    def fill(self, content_map: dict) -> None:
        """填充内容到文档"""
        # 处理段落
        self._fill_paragraphs(content_map)
        
        # 处理表格（增强版）
        self._fill_tables_enhanced(content_map)
        
        self.doc.save(self.output_path)
    
    def _fill_paragraphs(self, content_map: dict) -> None:
        """填充段落内容"""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            style = para.style.name
            
            # 找到标题段落
            if style.startswith('Heading'):
                title = text
                if title in content_map:
                    # 在标题后插入内容
                    self._insert_content_after_heading(para, content_map[title])
    
    def _insert_content_after_heading(self, heading_para, content: str) -> None:
        """在标题段落后插入内容"""
        if not heading_para.runs:
            heading_para.add_run()
        
        content_para = heading_para.insert_paragraph_before('')
        content_para.style = 'Normal'
        
        run = content_para.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        
        content_para.paragraph_format.line_spacing = Pt(20)
        content_para.paragraph_format.space_after = Pt(6)
    
    def _fill_tables_enhanced(self, content_map: dict) -> None:
        """增强版表格填充 - 智能识别字段并填充"""
        print("   处理表格字段...")
        
        # 收集所有字段的要求信息（用于 LLM 生成）
        field_requirements = {}
        
        # 个人信息字段列表（需要精确信息）
        personal_fields = [
            '学校名称', '学校全称', '团队名称', '联系电话', '邮箱',
            '指导教师', '指导教师姓名', '队长姓名', '队员姓名', '团队编号',
            '参赛赛项', '参赛赛项名称',
        ]
        
        for table_idx, table in enumerate(self.doc.tables):
            print(f"     表格 {table_idx + 1}/{len(self.doc.tables)}")
            
            # 解析表格字段
            fields = self.table_parser.parse_table(table)
            
            for field in fields:
                std_name = field['std_name']
                label = field['label']
                cell = field['cell']
                requirements = field['requirements']
                word_limit = field['word_limit']
                content_points = field['content_points']
                
                # 存储字段要求（用于后续 LLM 生成）
                if std_name not in field_requirements:
                    field_requirements[std_name] = {
                        'requirements': requirements,
                        'word_limit': word_limit,
                        'content_points': content_points,
                    }
                
                # 查找匹配的内容
                matched_content = self._find_matched_content(std_name, label, content_map)
                
                # 检查是否是个人信息字段
                is_personal = std_name in personal_fields or any(p in label for p in personal_fields)
                
                if matched_content:
                    # 检查置信度
                    confidence = self._get_confidence(std_name, label, content_map)
                    
                    # 个人信息字段：只有高置信度才填充
                    if is_personal and confidence < 0.7:
                        print(f"       ⚠️ {label} -> 跳过（个人信息，置信度{confidence:.0%}）")
                        continue
                    
                    # 填充内容
                    self._fill_table_cell(cell, matched_content, requirements, word_limit)
                    print(f"       ✅ {label} -> {len(matched_content)}字")
                elif is_personal:
                    # 个人信息字段无内容：不填充，保持空白
                    print(f"       ⚪ {label} -> 跳过（个人信息，用户未提供）")
        
        # 返回字段要求信息，供外部使用
        self.field_requirements = field_requirements
    
    def _find_matched_content(self, std_name: str, label: str, content_map: dict) -> str:
        """查找匹配的内容"""
        # 直接匹配标准名称
        if std_name in content_map:
            content = content_map[std_name]
            if content and content.strip():
                return content
        
        # 匹配原始标签
        if label in content_map:
            content = content_map[label]
            if content and content.strip():
                return content
        
        # 模糊匹配
        for key, value in content_map.items():
            if std_name in key or key in std_name:
                if value and value.strip():
                    return value
            if label in key or key in label:
                if value and value.strip():
                    return value
        
        return None
    
    def _get_confidence(self, std_name: str, label: str, content_map: dict) -> float:
        """获取匹配的置信度（如果可用）"""
        # 尝试从 content_map 中获取置信度信息
        # 这里简化处理，返回默认值
        # 实际应用中可以从 matcher 获取更详细的置信度信息
        return 0.7  # 默认置信度
    
    def _fill_table_cell(self, cell, content: str, requirements: str = '', word_limit: int = None) -> None:
        """填充表格单元格"""
        # 清空单元格内容
        cell.text = ''
        
        # 如果没有从参数获取字数限制，尝试从要求中提取
        if word_limit is None and requirements:
            word_limit = self._extract_word_limit(requirements)
        
        # 如果有字数限制且内容超出，截断内容
        if word_limit and len(content) > word_limit:
            content = content[:word_limit-3] + '...'
        
        # 添加新内容
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
    
    def _extract_word_limit(self, requirements: str) -> int:
        """从要求中提取字数限制"""
        match = re.search(r'不超过？(\d+)[字个]', requirements)
        return int(match.group(1)) if match else None
    
    def _fill_tables(self, content_map: dict) -> None:
        """填充表格字段（旧版，保留兼容性）"""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    # 识别占位符 {{field_name}}
                    if '{{' in text and '}}' in text:
                        placeholders = re.findall(r'{{(\w+)}}', text)
                        for placeholder in placeholders:
                            if placeholder in content_map:
                                text = text.replace(f'{{{{{placeholder}}}}}', content_map[placeholder])
                        cell.text = text
    
    def add_fill_report(self, report: str, confidence: dict) -> None:
        """在文档末尾添加填充报告"""
        self.doc.add_page_break()
        
        report_heading = self.doc.add_heading('📋 自动填充报告', level=1)
        report_heading.paragraph_format.space_before = Pt(24)
        report_heading.paragraph_format.space_after = Pt(12)
        
        intro = self.doc.add_paragraph()
        intro.add_run('本文档由 材料模板填写 Skill 自动填充生成。\n').bold = True
        intro.add_run('请仔细检查以下内容，特别是标记为 ⚠️ 和 ❌ 的部分。\n')
        
        legend = self.doc.add_paragraph()
        legend.add_run('图例：').bold = True
        legend.add_run(' ✅ 高置信度  ')
        legend.add_run('⚠️ 中置信度（建议检查）  ')
        legend.add_run('❌ 低置信度（需人工补充）')
        
        self.doc.add_paragraph('─' * 50)
        
        self.doc.add_heading('各模块填充情况', level=2)
        
        for section_title, conf in confidence.items():
            if conf >= 0.7:
                status = "✅"
            elif conf >= 0.4:
                status = "⚠️"
            else:
                status = "❌"
            
            p = self.doc.add_paragraph()
            p.add_run(f'{status} {section_title}').bold = True
            p.add_run(f' (置信度：{conf:.0%})')
        
        self.doc.add_heading('建议', level=2)
        suggestions = self.doc.add_paragraph()
        suggestions.add_run('1. 打开文档检查自动填充的内容\n')
        suggestions.add_run('2. 重点检查 ⚠️ 和 ❌ 标记的部分\n')
        suggestions.add_run('3. 补充缺失的内容\n')
        suggestions.add_run('4. 调整格式以符合要求\n')
        
        self.doc.save(self.output_path)
    
    def get_output_info(self) -> dict:
        """返回输出文件信息"""
        return {
            'output_path': self.output_path,
            'template_path': self.template_path,
            'file_size': os.path.getsize(self.output_path) if os.path.exists(self.output_path) else 0
        }


def main():
    import sys
    if len(sys.argv) < 3:
        print("用法：python docx_filler.py <模板文件> <输出文件>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    test_content = {
        '项目背景': '当前在线教育平台缺乏个性化推荐，学生学习效率低。',
        '研究目标': '开发一个智能学习推荐系统。',
        '研究内容': '包括用户行为分析、内容推荐算法。',
    }
    
    try:
        filler = DocxFiller(template_path, output_path)
        filler.fill(test_content)
        filler.add_fill_report({}, {})
        print(f"✅ 填充完成：{output_path}")
        print(f"📄 文件信息：{filler.get_output_info()}")
    except Exception as e:
        print(f"❌ 错误：{e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
