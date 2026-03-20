#!/usr/bin/env python3
"""
测试：填充产教融合 5G+ 创新应用赛模板
"""

import sys
import os
sys.path.insert(0, 'scripts')

from docx import Document
import shutil
from datetime import datetime

# 模板和输出
template_file = 'template_test.docx'
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
output_file = f'产教融合 5G+ 申报书_filled_{timestamp}.docx'

# 示例项目内容
project_data = {
    '参赛赛项名称': '5G+ 智慧教育',
    '作品/方案名称': '基于 5G+AI 的沉浸式远程教学系统',
    '学校全称': 'XX 大学',
    '团队名称': '智汇未来团队',
    '联系电话': '13800138000',
    '邮箱': 'team@example.edu.cn',
    '指导教师姓名': '张教授',
    '任课专业': '计算机科学与技术',
    '队长姓名': '李明',
    '所在专业': '软件工程',
    '摘要': '''本项目针对远程教学中互动性差、沉浸感不足的问题，提出基于 5G+AI 的沉浸式远程教学系统。
系统利用 5G 低延迟特性实现实时互动，结合 AI 技术实现智能课堂分析和个性化学习推荐。
核心成果包括：1）5G 超低延迟视频传输架构；2）基于计算机视觉的课堂行为分析系统；
3）AI 驱动的学习效果评估模型。已在 3 所学校试点，用户满意度达 92%。''',
    '设计目标': '''1. 实现端到端延迟<50ms 的远程互动教学
2. 课堂行为识别准确率>90%
3. 支持 1000+ 并发用户
4. 构建完整的沉浸式教学体验，包括虚拟实验室、AR 教学场景等''',
    '作品详情/解决方案详情': '''【系统架构】
系统采用三层架构：
- 接入层：5G CPE 设备 + 边缘计算节点
- 平台层：视频处理引擎、AI 分析引擎、数据中台
- 应用层：互动课堂、虚拟实验室、学情分析

【关键技术】
1. 5G 网络切片技术：保障教学视频 QoS
2. WebRTC+SFU 架构：支持大规模并发
3. YOLOv8+Transformer：课堂行为识别
4. 知识图谱：学习路径推荐

【创新点】
1. 首创 5G 边缘计算 + 云渲染的 AR 教学方案
2. 多模态学习分析（表情、姿态、语音）
3. 自适应学习路径生成

【部署方案】
- 学校端：5G CPE + 边缘服务器
- 云端：中心管理平台 + 资源库
- 终端：PC/平板/VR 头显''',
    '经济与社会价值': '''【商业价值】
- 目标市场：K12 学校、高校、职业培训机构
- 盈利模式：SaaS 订阅 + 硬件销售 + 定制开发
- 预计 3 年内服务 500 所学校，营收 5000 万

【社会价值】
- 促进教育公平，让偏远地区享受优质教育资源
- 提升教学效率 30% 以上
- 减少教师重复性工作，聚焦个性化指导

【竞品分析】
- 优势：5G 低延迟、AI 深度集成、沉浸式体验
- 差异化：边缘计算架构、多模态分析''',
    '项目进度计划': '''第一阶段（1-2 月）：需求调研与系统设计
  - 完成 10 所学校调研
  - 输出系统架构设计文档
  
第二阶段（3-5 月）：核心功能开发
  - 5G 视频传输引擎
  - AI 行为分析模块
  
第三阶段（6-8 月）：系统集成与测试
  - 端到端联调
  - 性能优化
  
第四阶段（9-10 月）：试点部署
  - 3 所学校试运行
  - 收集反馈并迭代
  
第五阶段（11-12 月）：验收与推广
  - 项目验收
  - 市场推广计划''',
}

# 复制模板
shutil.copy(template_file, output_file)
doc = Document(output_file)

print("=" * 60)
print("📝 开始填充模板：产教融合 5G+ 创新应用赛")
print("=" * 60)

# 填充表格
fill_count = 0
for table in doc.tables:
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            
            # 查找匹配的字段
            for field_name, field_value in project_data.items():
                if field_name in cell_text or cell_text in field_name:
                    # 找到对应单元格，在下一列或下方填写内容
                    if cell_idx + 1 < len(row.cells):
                        # 右侧单元格
                        next_cell = row.cells[cell_idx + 1]
                        if not next_cell.text.strip():
                            next_cell.text = field_value
                            print(f"✅ 填充：{field_name[:20]}...")
                            fill_count += 1
                    break

doc.save(output_file)

print("\n" + "=" * 60)
print(f"✅ 填充完成！")
print(f"📄 输出文件：{output_file}")
print(f"📊 填充字段数：{fill_count}")
print(f"📁 文件大小：{os.path.getsize(output_file) / 1024:.1f} KB")
print("=" * 60)

# 验证填充结果
print("\n🔍 验证填充结果：")
doc_check = Document(output_file)
filled_cells = []
for table in doc_check.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and any(proj in text for proj in project_data.values()):
                filled_cells.append(text[:50])

print(f"   已填充内容单元格：{len(filled_cells)} 个")
for i, cell in enumerate(filled_cells[:5], 1):
    print(f"   {i}. {cell}...")

print("\n💡 请用 Word 打开检查填充效果！")
