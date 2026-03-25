#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
材料模板填充 Web 应用 - Flask 后端服务器
完全使用 material-template-filler skill 的原有逻辑和目录结构

支持三种部署模式：
1. 独立部署：web/ 目录放在 skill 根目录下，使用相对路径
2. OpenClaw 部署：自动检测 ~/.openclaw/workspace
3. 自定义：通过环境变量指定路径
"""

import os
import sys
import json
import uuid
import shutil
import subprocess
import re
from datetime import datetime
from pathlib import Path
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename

app = Flask(__name__)

# 配置
PORT = int(os.environ.get('PORT', 5000))
MAX_UPLOAD_SIZE = int(os.environ.get('MAX_UPLOAD_SIZE', 50)) * 1024 * 1024

# ============ 路径配置（优先级：环境变量 > 相对路径 > OpenClaw 默认路径）============

def detect_paths():
    """自动检测并配置路径"""
    
    # 1. 优先使用环境变量
    if os.environ.get('WORKSPACE_DIR'):
        print("📍 使用环境变量配置的路径")
        workspace = Path(os.environ['WORKSPACE_DIR'])
        templates = Path(os.environ.get('TEMPLATES_DIR', workspace / 'templates'))
        inputs = Path(os.environ.get('INPUTS_DIR', workspace / 'inputs'))
        filled = Path(os.environ.get('FILLED_DIR', workspace / 'filled'))
        skill_scripts = Path(os.environ.get('SKILL_SCRIPTS_DIR', workspace / 'skills' / 'material-template-filler' / 'scripts'))
        return workspace, templates, inputs, filled, skill_scripts
    
    # 2. 尝试相对路径（web/ 在 skill 根目录下）
    web_dir = Path(__file__).parent
    skill_root = web_dir.parent
    if (skill_root / 'scripts' / 'main.py').exists():
        print("📍 使用相对路径（独立部署模式）")
        # 独立部署：在 skill 目录下创建 workspace
        workspace = skill_root / 'workspace'
        templates = workspace / 'templates'
        inputs = workspace / 'inputs'
        filled = workspace / 'filled'
        skill_scripts = skill_root / 'scripts'
        return workspace, templates, inputs, filled, skill_scripts
    
    # 3. 回退到 OpenClaw 默认路径
    print("📍 使用 OpenClaw 默认路径")
    workspace = Path.home() / '.openclaw' / 'workspace'
    templates = workspace / 'templates'
    inputs = workspace / 'inputs'
    filled = workspace / 'filled'
    skill_scripts = workspace / 'skills' / 'material-template-filler' / 'scripts'
    return workspace, templates, inputs, filled, skill_scripts


WORKSPACE, TEMPLATES_DIR, INPUTS_DIR, FILLED_DIR, SKILL_DIR = detect_paths()
SKILL_MAIN = SKILL_DIR / 'main.py'

# 确保目录存在
for dir_path in [TEMPLATES_DIR, INPUTS_DIR, FILLED_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# 验证 skill 主程序存在
if not SKILL_MAIN.exists():
    print(f"⚠️  警告：未找到 skill 主程序 {SKILL_MAIN}")
    print("   请确保 material-template-filler skill 已正确安装")

app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_SIZE

# 允许的文件扩展名
ALLOWED_TEMPLATE_EXTENSIONS = {'docx'}
ALLOWED_INPUT_EXTENSIONS = {'md', 'txt', 'docx'}


def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def generate_timestamp():
    return datetime.now().strftime('%Y%m%d_%H%M%S')


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/upload/template', methods=['POST'])
def upload_template():
    """上传模板文件到 workspace/templates/"""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if not allowed_file(file.filename, ALLOWED_TEMPLATE_EXTENSIONS):
        return jsonify({'error': '只支持 .docx 格式的模板文件'}), 400
    
    timestamp = generate_timestamp()
    safe_filename = secure_filename(file.filename)
    save_path = TEMPLATES_DIR / f"{timestamp}_{safe_filename}"
    
    file.save(save_path)
    
    return jsonify({
        'success': True,
        'filename': safe_filename,
        'path': str(save_path)
    })


@app.route('/api/upload/input', methods=['POST'])
def upload_input():
    """上传项目说明文件到 workspace/inputs/"""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if not allowed_file(file.filename, ALLOWED_INPUT_EXTENSIONS):
        return jsonify({'error': '只支持 .md, .txt, .docx 格式'}), 400
    
    timestamp = generate_timestamp()
    safe_filename = secure_filename(file.filename)
    name, ext = os.path.splitext(safe_filename)
    save_path = INPUTS_DIR / f"{name}_{timestamp}{ext}"
    
    file.save(save_path)
    
    return jsonify({
        'success': True,
        'filename': save_path.name,
        'path': str(save_path)
    })


@app.route('/api/list-inputs', methods=['GET'])
def list_inputs():
    """列出历史项目说明文件"""
    files = []
    for f in sorted(INPUTS_DIR.glob('*'), key=lambda x: x.stat().st_mtime, reverse=True):
        if f.is_file() and f.suffix in ['.md', '.txt', '.docx']:
            files.append({
                'name': f.name,
                'path': str(f),
                'size': f.stat().st_size,
                'mtime': datetime.fromtimestamp(f.stat().st_mtime).strftime('%Y-%m-%d %H:%M')
            })
    return jsonify({'files': files})


@app.route('/api/fill', methods=['POST'])
def fill_template():
    """执行模板填充 - 完全使用原始 skill 逻辑"""
    data = request.json
    
    template_path = data.get('template_path')
    input_path = data.get('input_path')  # 项目说明文件路径
    content = data.get('content', '')    # 或直接输入的内容
    generate_mindmap = data.get('generate_mindmap', False)  # 是否生成思维导图
    mermaid_content = data.get('mermaid_content', '')  # 自定义 Mermaid 内容
    
    if not template_path or not Path(template_path).exists():
        return jsonify({'error': '模板文件不存在'}), 400
    
    # 确定输入内容
    user_input = ''
    
    if input_path and Path(input_path).exists():
        # 使用上传的项目说明文件 - 读取文件内容
        try:
            ext = Path(input_path).suffix.lower()
            if ext in ['.md', '.txt']:
                with open(input_path, 'r', encoding='utf-8') as f:
                    user_input = f.read()
            elif ext == '.docx':
                # 读取 docx 文件内容
                from docx import Document
                doc = Document(input_path)
                user_input = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                return jsonify({'error': '不支持的文件格式'}), 400
        except Exception as e:
            return jsonify({'error': f'读取文件失败：{str(e)}'}), 500
    elif content.strip():
        # 使用直接输入的内容，保存到 inputs/
        timestamp = generate_timestamp()
        # 尝试提取项目名称作为文件名
        project_name = "项目说明"
        for line in content.split('\n'):
            if '项目名称' in line and ':' in line:
                project_name = line.split(':')[1].strip()
                break
        input_path = INPUTS_DIR / f"{project_name}_{timestamp}.md"
        with open(input_path, 'w', encoding='utf-8') as f:
            f.write(content)
        user_input = content
    else:
        return jsonify({'error': '请提供项目内容或上传项目说明文件'}), 400
    
    # 使用 docx_filler_v3 进行填充（支持图片插入）
    try:
        # 生成输出文件名
        timestamp = generate_timestamp()
        output_filename = f"filled_{timestamp}.docx"
        output_path = FILLED_DIR / output_filename
        
        # 导入 docx_filler_v3
        sys.path.insert(0, str(SKILL_DIR))
        from docx_filler_v3 import DocxFillerV3
        
        # 创建填充器
        filler = DocxFillerV3(template_path, str(output_path))
        
        # 填充文本内容
        filler.fill(user_input)
        
        # 生成并插入思维导图
        mindmap_generated = False
        if generate_mindmap:
            if mermaid_content.strip():
                # 使用自定义 Mermaid 内容
                success = filler.insert_mindmap(
                    mermaid_content=mermaid_content,
                    caption='图 1: 项目思维导图',
                    method='dot'
                )
                mindmap_generated = success
            else:
                # 从文本自动生成（简单实现：提取关键词）
                mindmap_content = generate_mermaid_from_text(user_input)
                if mindmap_content:
                    success = filler.insert_mindmap(
                        mermaid_content=mindmap_content,
                        caption='图 1: 项目思维导图',
                        method='dot'
                    )
                    mindmap_generated = success
        
        # 添加填充报告
        filler.add_fill_report()
        
        # 保存文档
        filler.doc.save(str(output_path))
        
        # 获取输出文件信息
        output_file = output_path
        report_content = ''  # 简化实现，暂不生成报告
        
        return jsonify({
            'success': True,
            'output_filename': output_file.name,
            'report_filename': None,
            'report_content': report_content,
            'log': ''
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/<filename>')
def download_file(filename):
    """从 filled/ 目录下载文件"""
    file_path = FILLED_DIR / filename
    
    if not file_path.exists():
        return jsonify({'error': '文件不存在'}), 404
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=filename
    )


@app.route('/api/generate-mindmap', methods=['POST'])
def generate_mindmap_api():
    """从文本生成思维导图 Mermaid 代码"""
    data = request.json
    text = data.get('text', '')
    
    if not text.strip():
        return jsonify({'error': '请提供文本内容'}), 400
    
    try:
        mermaid = generate_mermaid_from_text(text)
        return jsonify({
            'success': True,
            'mermaid': mermaid
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/read-file', methods=['POST'])
def read_file_api():
    """读取文件内容（用于思维导图生成）"""
    data = request.json
    file_path = data.get('path', '')
    
    if not file_path or not Path(file_path).exists():
        return jsonify({'error': '文件不存在'}), 400
    
    try:
        ext = Path(file_path).suffix.lower()
        content = ''
        
        if ext in ['.md', '.txt']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            return jsonify({'error': '不支持的文件格式'}), 400
        
        return jsonify({
            'success': True,
            'content': content
        })
    except Exception as e:
        return jsonify({'error': f'读取文件失败：{str(e)}'}), 500


@app.route('/api/chat-improve', methods=['POST'])
def chat_improve():
    """对话式改进文档"""
    data = request.json
    message = data.get('message', '')
    last_output = data.get('last_output', '')
    original_content = data.get('original_content', '')
    
    if not message:
        return jsonify({'error': '请提供修改要求'}), 400
    
    try:
        # 生成改进建议的回复
        improvement_response = generate_improvement_response(message, original_content)
        
        # 如果需要生成新版本文档
        new_version_file = None
        version = None
        
        if should_regenerate_document(message):
            # 生成新版本文档
            new_version_file = regenerate_document(last_output, message, original_content)
            if new_version_file:
                version = 'v2'
        
        return jsonify({
            'success': True,
            'response': improvement_response,
            'new_version_file': new_version_file,
            'version': version
        })
        
    except Exception as e:
        return jsonify({'error': f'处理失败：{str(e)}'}), 500


def generate_improvement_response(message: str, content: str) -> str:
    """
    生成改进建议的回复（简化实现）
    """
    message_lower = message.lower()
    
    if '详细' in message_lower or '更多' in message_lower:
        return "好的，我会让内容更加详细。建议在以下几个部分增加更多细节：\n\n1. **项目背景**：可以补充行业现状、市场规模等数据\n2. **技术方案**：可以详细说明技术路线、实现细节\n3. **创新点**：可以具体描述与现有方案的差异\n\n请告诉我具体想详细哪个部分，我会帮你生成新版本文档。"
    
    elif '创新' in message_lower:
        return "好的，我来帮你强化创新点的描述。建议从以下角度突出创新性：\n\n1. **技术创新**：使用了什么新技术或新方法\n2. **应用创新**：解决了什么之前没解决的问题\n3. **模式创新**：有什么新的商业模式或应用场景\n\n需要我帮你生成一个强调创新点的新版本吗？"
    
    elif '正式' in message_lower or '语气' in message_lower:
        return "好的，我会调整文档语气，使其更加正式和专业。主要会：\n\n1. 使用更正式的词汇和表达\n2. 增加客观数据支撑\n3. 减少口语化表达\n4. 使用更规范的文档格式\n\n需要我现在生成一个正式版本吗？"
    
    elif '简短' in message_lower or '简洁' in message_lower:
        return "好的，我会让内容更加简洁明了。主要会：\n\n1. 删除冗余描述\n2. 突出核心要点\n3. 使用更精炼的语言\n4. 保留关键数据和结论\n\n需要我现在生成一个简洁版本吗？"
    
    else:
        return f"收到你的修改要求：\"{message}\"\n\n我可以帮你：\n• 调整内容详细程度\n• 强化特定部分（如创新点、技术细节）\n• 调整文档语气和风格\n• 优化结构和逻辑\n\n请告诉我具体想怎么改，我会生成新版本文档。"


def should_regenerate_document(message: str) -> bool:
    """判断是否需要重新生成文档"""
    keywords = ['生成', '修改', '调整', '重写', '改进', '优化', '版本', '详细', '简洁', '正式']
    return any(keyword in message for keyword in keywords)


def regenerate_document(last_output: str, message: str, original_content: str) -> str:
    """
    重新生成文档（简化实现：复制原文档）
    实际应该调用 AI API 生成新内容
    """
    if not last_output:
        return None
    
    import os
    from datetime import datetime
    import shutil
    
    base_name = os.path.splitext(last_output)[0]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    new_filename = f"{base_name}_v2_{timestamp}.docx"
    new_path = FILLED_DIR / new_filename
    
    # 复制原文档作为新版本
    last_path = FILLED_DIR / last_output
    
    if last_path.exists():
        shutil.copy(last_path, new_path)
        return new_filename
    
    return None


def generate_mermaid_from_text(text: str) -> str:
    """
    从文本生成简单的 Mermaid 思维导图
    
    这是一个简化实现，提取文本中的关键词和结构
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # 提取项目名称作为根节点
    root_name = "项目主题"
    for line in lines[:5]:  # 在前 5 行查找
        if '项目名称' in line:
            if ':' in line:
                root_name = line.split(':')[1].strip()
            else:
                root_name = line.replace('项目名称', '').strip()
            break
    
    # 构建思维导图
    mermaid_lines = [f"mindmap", f"  root(({root_name}))"]
    
    # 提取关键章节
    sections = []
    current_section = None
    
    for line in lines:
        # 检查是否是章节标题
        if re.match(r'^[一二三四五六七八九十]+[、.．]', line) or \
           re.match(r'^\d+[.．]', line) or \
           re.match(r'^第 [一二三四五六七八九十\d]+[章节]', line):
            # 提取章节名称
            section_name = re.sub(r'^[一二三四五六七八九十\d]+[、.．第章节]', '', line).strip()
            if section_name and len(section_name) < 50:
                current_section = section_name
                sections.append(current_section)
                mermaid_lines.append(f"    {current_section}")
        elif current_section and len(line) < 100:
            # 提取子节点（短行）
            if ':' in line:
                key = line.split(':')[0].strip()
                if key and len(key) < 30:
                    mermaid_lines.append(f"      {key}")
    
    # 如果没有提取到章节，使用默认结构
    if len(mermaid_lines) <= 2:
        mermaid_lines = [
            "mindmap",
            f"  root(({root_name}))",
            "    项目概述",
            "      背景",
            "      目标",
            "    技术方案",
            "      核心功能",
            "      创新点",
            "    实施计划",
            "      阶段划分",
            "      时间安排"
        ]
    
    return '\n'.join(mermaid_lines)


if __name__ == '__main__':
    print(f"🚀 材料模板填充 Web 应用启动中...")
    print(f"📂 工作目录：{WORKSPACE}")
    print(f"📁 模板目录：{TEMPLATES_DIR}")
    print(f"📁 项目说明目录：{INPUTS_DIR}")
    print(f"📁 输出目录：{FILLED_DIR}")
    print(f"🌐 访问地址：http://localhost:{PORT}")
    app.run(host='0.0.0.0', port=PORT, debug=False)
